from __future__ import annotations

import shutil
from pathlib import Path

from scripts.workspace import init_workspace


def write_discussion_to_docx(project: Path | str, manuscript: Path | str, output: Path | str | None = None, heading: str = "Discussion") -> Path:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise ValueError("python-docx is required for DOCX writeback") from exc
    root = Path(project).resolve()
    ws = init_workspace(root)
    final = ws / "discussion_final.md"
    if not final.exists() or not final.read_text(encoding="utf-8").strip():
        raise ValueError("discussion_final.md is missing or empty; pass the release gate first")
    source = Path(manuscript).resolve()
    destination = Path(output).resolve() if output else ws / f"{source.stem}-with-discussion.docx"
    shutil.copy2(source, destination)
    document = Document(destination)
    document.add_heading(heading, level=1)
    for paragraph in [x.strip() for x in final.read_text(encoding="utf-8").split("\n\n") if x.strip()]:
        document.add_paragraph(paragraph)
    document.save(destination)
    return destination
