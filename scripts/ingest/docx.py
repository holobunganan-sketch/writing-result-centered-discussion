from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from .base import ExtractedUnit, unit


def _fallback(path: Path) -> list[ExtractedUnit]:
    units: list[ExtractedUnit] = []
    with zipfile.ZipFile(path) as archive:
        if "word/document.xml" not in archive.namelist():
            return []
        root = ET.fromstring(archive.read("word/document.xml"))
        p_no = 0
        for elem in root.iter():
            if elem.tag.rsplit("}", 1)[-1] == "p":
                text = "".join(x.text or "" for x in elem.iter() if x.tag.rsplit("}", 1)[-1] == "t")
                if text.strip():
                    p_no += 1
                    if found := unit(f"paragraph {p_no}", text, format="docx"):
                        units.append(found)
    return units


def extract_docx(path: Path) -> list[ExtractedUnit]:
    try:
        from docx import Document  # type: ignore
        document = Document(path)
        units: list[ExtractedUnit] = []
        p_no = 0
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                p_no += 1
                if found := unit(f"paragraph {p_no}", paragraph.text, format="docx", style=paragraph.style.name if paragraph.style else ""):
                    units.append(found)
        for t_idx, table in enumerate(document.tables, 1):
            for r_idx, row in enumerate(table.rows, 1):
                for c_idx, cell in enumerate(row.cells, 1):
                    if found := unit(f"table {t_idx} row {r_idx} col {c_idx}", cell.text, format="docx-table"):
                        units.append(found)
        for section_idx, section in enumerate(document.sections, 1):
            for kind, collection in (("header", section.header.paragraphs), ("footer", section.footer.paragraphs)):
                for p_idx, paragraph in enumerate(collection, 1):
                    if found := unit(f"section {section_idx} {kind} paragraph {p_idx}", paragraph.text, format="docx"):
                        units.append(found)
        return units or _fallback(path)
    except Exception:
        return _fallback(path)
